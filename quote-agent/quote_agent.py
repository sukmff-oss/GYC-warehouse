"""
國研化工智能報價 Agent
自動生成格式化的 Word 報價單

用法：
python quote_agent.py --customer "三芳" --products "WP-532P:20,WP-730P:20" --contact "許清秀"
"""

import sys
import os
import argparse
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(__file__))
from product_db import PRODUCTS, CUSTOMERS, GY_INFO


def find_product(model_code):
    """模糊搜尋產品"""
    model_code = model_code.upper().strip()
    for p in PRODUCTS:
        if p["型號"].upper() == model_code:
            return p
    # 模糊比對
    for p in PRODUCTS:
        if model_code in p["型號"].upper():
            return p
    return None


def find_customer(name):
    """模糊搜尋客戶"""
    name = name.strip()
    for key in CUSTOMERS:
        if name in key or name in CUSTOMERS[key]["簡稱"]:
            return key, CUSTOMERS[key]
    return None, None


def create_quote_document(
    customer_name,
    contact,
    products,  # list of (model_code, quantity_kg)
    valid_days=30,
    notes="",
    output_path=None
):
    """產生報價單 Word 文件"""
    doc = Document()

    # 頁面設定：橫向
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ===== 公司抬頭 =====
    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run(GY_INFO["公司名稱"])
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "微軟正黑體"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    # 報價單標題
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("報  價  單")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "微軟正黑體"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    run = title.add_run("  E S T I M A T E")
    run.font.size = Pt(10)
    run.font.name = "Arial"

    # ===== 客戶資料區 =====
    doc.add_paragraph()  # 空行

    cust_key, cust_data = find_customer(customer_name)
    if cust_key:
        display_name = cust_key
    else:
        display_name = customer_name
        cust_data = {}

    # ATTN 那一行
    attn_line = doc.add_paragraph()
    attn_line.paragraph_format.space_after = Pt(2)
    run = attn_line.add_run("ATTN : ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微軟正黑體"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    run = attn_line.add_run(contact if contact else (cust_data.get("主要聯絡人", [""])[0] if cust_data.get("主要聯絡人") else ""))
    run.font.size = Pt(11)
    run.font.name = "微軟正黑體"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    customer_line = doc.add_paragraph()
    customer_line.paragraph_format.space_after = Pt(2)
    run = customer_line.add_run(f"{display_name} 台照")
    run.font.size = Pt(11)
    run.font.name = "微軟正黑體"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    # 日期
    now = datetime.now()
    ROC_YEAR = now.year - 1911
    date_line = doc.add_paragraph()
    run = date_line.add_run(f"中華民國 {ROC_YEAR} 年 {now.month} 月 {now.day} 日  即日起生效")
    run.font.size = Pt(11)
    run.font.name = "微軟正黑體"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    run = date_line.add_run(f"　（有效期限：{valid_days}天）")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

    # ===== 產品表格 =====
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # 欄位標題
    hdr = table.rows[0].cells
    headers = ["品　　名", "包　　裝", "單　　價", "備　　註"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr[i].paragraphs[0].runs[0].font.size = Pt(11)
        hdr[i].paragraphs[0].runs[0].font.name = "微軟正黑體"
        hdr[i].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    total = 0
    for model_code, qty_kg in products:
        p = find_product(model_code)
        row = table.add_row().cells

        if p:
            name_text = f"{p['型號']} {p['品名']}"
            package = p["包裝"] or ""
            price = p["單價"]
            remark = p["備註"] or ""
        else:
            name_text = f"{model_code}（待確認）"
            package = "待確認"
            price = None
            remark = "產品資料庫中無此型號"

        row[0].text = name_text
        row[1].text = package
        row[2].text = f"{price:,}" if price else "待確認"
        row[3].text = remark

        for cell in row:
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.name = "微軟正黑體"
            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not price and "待確認" in cell.text:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        if price:
            total += price * qty_kg

    # 總價（估算）
    if total > 0:
        note_row = table.add_row().cells
        note_row[0].merge(note_row[3])
        note_row[0].text = f"※ 以上報價為 NTD/KG，實際金額依訂單數量與規格調整。粗估參考總值：NT$ {total:,}"
        note_row[0].paragraphs[0].runs[0].font.size = Pt(10)
        note_row[0].paragraphs[0].runs[0].italic = True
        note_row[0].paragraphs[0].runs[0].font.name = "微軟正黑體"
        note_row[0].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    # ===== 備註 =====
    if notes:
        doc.add_paragraph()
        note_p = doc.add_paragraph()
        run = note_p.add_run(f"備註：{notes}")
        run.font.size = Pt(10)
        run.font.name = "微軟正黑體"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    # ===== 公司資訊 =====
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lines = [
        f"大 園 廠：{GY_INFO['工廠地址']}",
        f"電  話：{GY_INFO['電話']}",
        f"傳  真：{GY_INFO['傳真']}",
        f"統一編號：{GY_INFO['統一編號']}",
        f"網  址：{GY_INFO['網址']}",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.name = "微軟正黑體"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    # ===== 存檔 =====
    if not output_path:
        filename = f"報價單_{display_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = os.path.join(os.path.dirname(__file__), filename)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="國研化工智能報價 Agent")
    parser.add_argument("--customer", "-c", required=True, help="客戶名稱")
    parser.add_argument("--products", "-p", required=True, help="產品列表，格式：型號:數量,型號:數量")
    parser.add_argument("--contact", help="聯絡人")
    parser.add_argument("--valid", type=int, default=30, help="有效天數")
    parser.add_argument("--notes", help="備註")
    parser.add_argument("--output", "-o", help="輸出檔名")

    args = parser.parse_args()

    # 解析產品
    products = []
    for item in args.products.split(","):
        item = item.strip()
        if ":" in item:
            code, qty = item.split(":")
            products.append((code.strip(), int(qty.strip())))
        else:
            products.append((item, 20))  # 預設20KG

    output = create_quote_document(
        customer_name=args.customer,
        contact=args.contact,
        products=products,
        valid_days=args.valid,
        notes=args.notes or "",
        output_path=args.output
    )
    print(f"✅ 報價單已生成：{output}")