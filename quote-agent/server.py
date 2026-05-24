"""
國研化工智能報價系統 - Flask 後端
啟動方式：python server.py
"""

from flask import Flask, request, send_file, render_template
from flask_cors import CORS
import sys, os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(__file__))
from product_db import PRODUCTS, CUSTOMERS, GY_INFO

app = Flask(__name__, template_folder='.')
CORS(app)

def find_product(model_code):
    model_code = model_code.upper().strip()
    for p in PRODUCTS:
        if p["型號"].upper() == model_code:
            return p
    for p in PRODUCTS:
        if model_code in p["型號"].upper():
            return p
    return None

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/generate", methods=["POST"])
def generate():
    data = request.get_json()
    customer = data.get("customer", "")
    contact = data.get("contact", "")
    valid_days = data.get("valid_days", 30)
    notes = data.get("notes", "")
    products = data.get("products", [])

    doc = Document()

    # 橫向
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    def add_run(para, text, bold=False, size=11, color=None):
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "微軟正黑體"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
        if color:
            run.font.color.rgb = RGBColor(*color)
        return run

    # 公司抬頭
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, GY_INFO["公司名稱"], bold=True, size=18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "報  價  單", bold=True, size=16)
    add_run(p, "  E S T I M A T E", size=10)

    doc.add_paragraph()

    # ATTN
    p = doc.add_paragraph()
    add_run(p, "ATTN : ", bold=True, size=11)
    add_run(p, contact, size=11)

    p = doc.add_paragraph()
    add_run(p, f"{customer} 台照", size=11)

    now = datetime.now()
    roc_year = now.year - 1911
    p = doc.add_paragraph()
    add_run(p, f"中華民國 {roc_year} 年 {now.month} 月 {now.day} 日  即日起生效")
    add_run(p, f"　（有效期限：{valid_days}天）", size=9, color=(0x99, 0x00, 0x00))

    doc.add_paragraph()

    # 表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    for i, h in enumerate(["品　　名", "包　　裝", "單　　價", "備　　註"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr[i].paragraphs[0].runs[0].font.size = Pt(11)

    total = 0
    for item in products:
        model = item.get("model", "")
        qty = item.get("qty", 0)
        p = find_product(model)
        row = table.add_row().cells

        if p:
            name_text = f"{p['型號']} {p['品名']}"
            package = p["包裝"] or ""
            price = p["單價"]
            remark = p["備註"] or ""
        else:
            name_text = f"{model}（待確認）"
            package = "待確認"
            price = None
            remark = "無此型號"

        row[0].text = name_text
        row[1].text = package
        row[2].text = f"{price:,}" if price else "待確認"
        row[3].text = remark

        for cell in row:
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if price:
            total += price * qty

    if total > 0:
        note_row = table.add_row().cells
        note_row[0].merge(note_row[3])
        note_row[0].text = f"※ 以上報價為 NTD/KG，實際金額依訂單數量與規格調整。粗估參考總值：NT$ {total:,}"
        note_row[0].paragraphs[0].runs[0].font.size = Pt(10)
        note_row[0].paragraphs[0].runs[0].italic = True

    if notes:
        doc.add_paragraph()
        p = doc.add_paragraph()
        add_run(p, f"備註：{notes}", size=10)

    doc.add_paragraph()
    for line in [
        f"大 園 廠：{GY_INFO['工廠地址']}",
        f"電  話：{GY_INFO['電話']}",
        f"傳  真：{GY_INFO['傳真']}",
        f"統一編號：{GY_INFO['統一編號']}",
        f"網  址：{GY_INFO['網址']}",
    ]:
        p = doc.add_paragraph()
        add_run(p, line, size=9)

    # 暫存
    filename = f"報價單_{customer}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(os.path.dirname(__file__), filename)
    doc.save(filepath)

    return send_file(filepath, as_attachment=True, download_name=filename)

if __name__ == "__main__":
    print("🧪 國研化工智能報價系統")
    print("📍 開啟瀏覽器前往：http://localhost:5000")
    app.run(host="0.0.0.0", port=5000, debug=True)